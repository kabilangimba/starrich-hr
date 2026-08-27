"""Text extraction from candidate CVs.

Supports PDF and DOCX (proposal §2). Extraction runs before any AI call, so a
malformed or image-only file fails here with a clear message rather than burning
an API call on empty input.
"""

from __future__ import annotations

import hashlib
import os

import frappe
from frappe import _

SUPPORTED_EXTENSIONS = (".pdf", ".docx")

#: Below this, treat the document as having no usable text layer - almost always
#: a scanned image saved as a PDF, which needs OCR rather than a parser.
MIN_USABLE_CHARS = 50


class UnsupportedResumeError(frappe.ValidationError):
	pass


def resolve_file_path(file_url: str) -> str:
	"""Map a Frappe file URL to an absolute path on disk.

	Goes through the File doctype rather than string-building a path, so private
	files resolve correctly and permissions are respected.
	"""
	if not file_url:
		raise UnsupportedResumeError(_("No resume file was provided."))

	name = frappe.db.get_value("File", {"file_url": file_url}, "name")
	if not name:
		raise UnsupportedResumeError(_("Could not find an uploaded file for {0}.").format(file_url))

	path = frappe.get_doc("File", name).get_full_path()
	if not os.path.exists(path):
		raise UnsupportedResumeError(_("The stored file is missing from disk: {0}.").format(file_url))

	return path


def extract_text(file_url: str) -> str:
	"""Return the plain text of a CV.

	Raises `UnsupportedResumeError` for unsupported types or files with no text
	layer, so the caller can tell the recruiter what to do about it.
	"""
	path = resolve_file_path(file_url)
	extension = os.path.splitext(path)[1].lower()

	if extension == ".pdf":
		text = _extract_pdf(path)
	elif extension == ".docx":
		text = _extract_docx(path)
	else:
		raise UnsupportedResumeError(
			_("{0} files are not supported. Upload a PDF or DOCX.").format(extension or _("Unknown"))
		)

	text = _normalise(text)

	if len(text) < MIN_USABLE_CHARS:
		raise UnsupportedResumeError(
			_(
				"No readable text was found in this file. If it is a scanned document, "
				"it needs to be run through OCR before it can be parsed."
			)
		)

	return text


def _extract_pdf(path: str) -> str:
	"""Extract with pdfplumber, falling back to pypdf.

	pdfplumber preserves layout better on multi-column CVs; pypdf is more tolerant
	of slightly malformed files, so trying both meaningfully improves the hit rate.
	"""
	try:
		import pdfplumber

		with pdfplumber.open(path) as pdf:
			pages = [page.extract_text() or "" for page in pdf.pages]
		text = "\n".join(pages)
		if len(_normalise(text)) >= MIN_USABLE_CHARS:
			return text
	except Exception:
		frappe.log_error(title="AI HR: pdfplumber extraction failed", message=frappe.get_traceback())

	try:
		from pypdf import PdfReader

		reader = PdfReader(path)
		return "\n".join((page.extract_text() or "") for page in reader.pages)
	except Exception:
		frappe.log_error(title="AI HR: pypdf extraction failed", message=frappe.get_traceback())
		raise UnsupportedResumeError(_("This PDF could not be read. It may be corrupt or encrypted."))


def _extract_docx(path: str) -> str:
	"""Extract paragraphs and table cells from a .docx.

	Many CVs lay out skills and dates in tables, whose text is not in the
	paragraph stream, so both are collected.
	"""
	try:
		import docx
	except ImportError:
		raise UnsupportedResumeError(
			_("The 'python-docx' package is not installed in this bench environment.")
		)

	try:
		document = docx.Document(path)
	except Exception:
		raise UnsupportedResumeError(
			_("This DOCX file could not be read. Legacy .doc files are not supported.")
		)

	parts = [p.text for p in document.paragraphs]
	for table in document.tables:
		for row in table.rows:
			parts.extend(cell.text for cell in row.cells)

	return "\n".join(parts)


def _normalise(text: str) -> str:
	"""Collapse whitespace so hashing is stable and token use stays low."""
	if not text:
		return ""
	lines = (line.strip() for line in text.splitlines())
	return "\n".join(line for line in lines if line).strip()


def content_hash(text: str) -> str:
	"""Stable fingerprint of CV text.

	Drives the re-parse guard in §17: identical text is never sent to the provider
	twice, which is the single biggest cost saving in the pipeline.
	"""
	return hashlib.sha256(text.encode("utf-8")).hexdigest()
